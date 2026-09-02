"""Render career/CV.md to CV.docx, CV.html and a one-page CV.pdf.

CV.md uses a deliberately small markdown subset so this renderer stays dumb:

    # Name                        -> name line (16pt bold)
    plain line straight after it  -> contact line
    ## Section                    -> section heading (11pt bold + bottom rule)
    **Lead** rest of line         -> entry/skill line (lead bold, rest regular)
    - text                        -> bullet (0.4cm hanging indent)
    anything else                 -> body paragraph

Everything from the `## Gaps` heading onwards is editorial notes for Samrath and
is never rendered.

PDF: Word COM if pywin32 is available, else Microsoft Edge headless printing
CV.html. Either way the result is asserted to be exactly one page.

Usage:  python career/build_cv.py [--src PATH] [--out DIR]
"""

from __future__ import annotations

import argparse
import html
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DEFAULT_SRC = Path(
    r"C:\Users\samra\OneDrive\dev\repos\ai-mastery\career\CV.md"
)
DEFAULT_OUT = DEFAULT_SRC.parent

EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

STOP_SECTION = "Gaps"  # section (and everything after) that never renders


# --------------------------------------------------------------------------
# parsing
# --------------------------------------------------------------------------

def parse(md: str) -> dict:
    """Turn the markdown subset into {name, contact, sections:[(title, blocks)]}."""
    name = ""
    contact = ""
    sections: list[tuple[str, list[tuple[str, str]]]] = []
    current: list[tuple[str, str]] | None = None

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            name = line[2:].strip()
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            if title.split(" ")[0] == STOP_SECTION:
                break
            current = []
            sections.append((title, current))
            continue
        if current is None:
            if not contact:
                contact = line.strip()
            continue
        if line.startswith("- "):
            current.append(("bullet", line[2:].strip()))
        elif line.startswith("**"):
            current.append(("entry", line.strip()))
        else:
            current.append(("body", line.strip()))

    if not name:
        raise SystemExit("CV.md has no `# Name` line")
    if not contact:
        raise SystemExit("CV.md has no contact line under the name")
    return {"name": name, "contact": contact, "sections": sections}


def split_bold(text: str) -> list[tuple[str, bool]]:
    """Split on **...** into (chunk, is_bold) pairs."""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    return [(chunk, bool(i % 2)) for i, chunk in enumerate(parts) if chunk]


# --------------------------------------------------------------------------
# docx
# --------------------------------------------------------------------------

def _bottom_rule(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 0.75pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    pPr.append(borders)


def _para(doc, *, size, bold=False, before=0, after=0, align=None,
          left=None, hanging=None, colour=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    if left is not None:
        pf.left_indent = left
    if hanging is not None:
        pf.first_line_indent = -hanging
    p._cv_style = (size, bold, colour)
    return p


def _run(p, text, *, size, bold, colour=None):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    if colour is not None:
        r.font.color.rgb = colour
    return r


def build_docx(cv: dict, path: Path) -> Path:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    p = _para(doc, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    _run(p, cv["name"], size=16, bold=True)

    p = _para(doc, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _run(p, cv["contact"], size=9, bold=False, colour=RGBColor(0x44, 0x44, 0x44))

    for title, blocks in cv["sections"]:
        head = _para(doc, size=11, before=7, after=3)
        _run(head, title.upper(), size=11, bold=True)
        _bottom_rule(head)

        for kind, text in blocks:
            if kind == "bullet":
                p = _para(doc, size=10, after=1,
                          left=Cm(0.4), hanging=Cm(0.4))
                _run(p, "\u2022  ", size=10, bold=False)
                for chunk, bold in split_bold(text):
                    _run(p, chunk, size=10, bold=bold)
            else:
                p = _para(doc, size=10, before=2 if kind == "entry" else 0,
                          after=1)
                for chunk, bold in split_bold(text):
                    _run(p, chunk, size=10, bold=bold)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


# --------------------------------------------------------------------------
# html (Edge fallback source, and useful on its own)
# --------------------------------------------------------------------------

def _inline_html(text: str) -> str:
    out = []
    for chunk, bold in split_bold(text):
        esc = html.escape(chunk)
        out.append(f"<strong>{esc}</strong>" if bold else esc)
    return "".join(out)


def build_html(cv: dict, path: Path) -> Path:
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        f"<title>{html.escape(cv['name'])} CV</title><style>",
        "@page { size: A4; margin: 1.5cm; }",
        "body { font-family: Calibri, 'Segoe UI', sans-serif; font-size: 10pt;",
        "       line-height: 1.18; margin: 0; color: #000; }",
        "h1 { font-size: 16pt; text-align: center; margin: 0 0 1pt; }",
        ".contact { font-size: 9pt; text-align: center; color: #444;",
        "           margin: 0 0 2pt; }",
        "h2 { font-size: 11pt; text-transform: uppercase; margin: 7pt 0 3pt;",
        "     border-bottom: 0.75pt solid #999; padding-bottom: 1pt; }",
        "p { margin: 0 0 1pt; }",
        "p.entry { margin-top: 2pt; }",
        "p.bullet { margin-left: 0.4cm; text-indent: -0.4cm; }",
        "</style></head><body>",
        f"<h1>{html.escape(cv['name'])}</h1>",
        f"<p class='contact'>{html.escape(cv['contact'])}</p>",
    ]
    for title, blocks in cv["sections"]:
        parts.append(f"<h2>{html.escape(title)}</h2>")
        for kind, text in blocks:
            cls = {"bullet": "bullet", "entry": "entry", "body": "body"}[kind]
            body = _inline_html(text)
            if kind == "bullet":
                body = "\u2022&nbsp;&nbsp;" + body
            parts.append(f"<p class='{cls}'>{body}</p>")
    parts.append("</body></html>")
    path.write_text("\n".join(parts), encoding="utf-8")
    return path


# --------------------------------------------------------------------------
# pdf
# --------------------------------------------------------------------------

def pdf_via_word(docx_path: Path, pdf_path: Path) -> int | None:
    """Export with Word. Returns Word's own page count, or None if unavailable."""
    try:
        import win32com.client.dynamic  # noqa: F401
    except ImportError:
        return None

    import win32com.client.dynamic as dyn

    word = None
    doc = None
    try:
        word = dyn.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path), False, True)
        doc.ExportAsFixedFormat(str(pdf_path), 17)  # 17 = wdExportFormatPDF
        return int(doc.ComputeStatistics(2))        # 2 = wdStatisticPages
    except Exception as exc:  # Word missing, broken COM cache, licence prompt
        print(f"word export unavailable: {exc}", file=sys.stderr)
        return None
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def pdf_via_edge(html_path: Path, pdf_path: Path) -> bool:
    if not EDGE.exists():
        print(f"edge not found at {EDGE}", file=sys.stderr)
        return False
    cmd = [
        str(EDGE),
        "--headless",
        "--disable-gpu",
        "--no-pdf-header-footer",
        f"--print-to-pdf={pdf_path}",
        html_path.as_uri(),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
    if not pdf_path.exists():
        print(proc.stderr.strip()[-800:], file=sys.stderr)
        return False
    return True


def count_pdf_pages(pdf_path: Path) -> int:
    data = pdf_path.read_bytes()
    n = len(re.findall(rb"/Type\s*/Page(?![s/\w])", data))
    if n:
        return n
    # some producers only declare the count on the page tree
    counts = [int(m) for m in re.findall(rb"/Count\s+(\d+)", data)]
    return max(counts) if counts else 0


# --------------------------------------------------------------------------

def main() -> int:
    ap = argparse.ArgumentParser(description="Render CV.md to docx/html/pdf.")
    ap.add_argument("--src", type=Path, default=DEFAULT_SRC)
    ap.add_argument("--out", type=Path, default=None)
    args = ap.parse_args()

    src: Path = args.src.resolve()
    if not src.exists():
        raise SystemExit(f"source not found: {src}")
    out: Path = (args.out or DEFAULT_OUT).resolve()
    out.mkdir(parents=True, exist_ok=True)

    cv = parse(src.read_text(encoding="utf-8"))
    words = sum(
        len(text.split())
        for _, blocks in cv["sections"]
        for _, text in blocks
    ) + len(cv["name"].split()) + len(cv["contact"].split())

    docx_path = build_docx(cv, out / "CV.docx")
    html_path = build_html(cv, out / "CV.html")
    pdf_path = out / "CV.pdf"
    if pdf_path.exists():
        pdf_path.unlink()

    pages = pdf_via_word(docx_path, pdf_path)
    producer = "word"
    if pages is None or not pdf_path.exists():
        producer = "edge"
        if not pdf_via_edge(html_path, pdf_path):
            raise SystemExit("no PDF producer available (Word COM and Edge both failed)")
        pages = count_pdf_pages(pdf_path)

    print(f"words={words}")
    print(f"docx={docx_path}")
    print(f"html={html_path}")
    print(f"pdf={pdf_path} (via {producer})")
    print(f"pages={pages}")

    if pages != 1:
        longest = max(
            cv["sections"],
            key=lambda s: sum(len(t.split()) for _, t in s[1]),
        )
        spill = sum(len(t.split()) for _, t in longest[1])
        print(
            f"CV is {pages} pages. Trim '{longest[0]}' first "
            f"({spill} words, the largest section).",
            file=sys.stderr,
        )
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
